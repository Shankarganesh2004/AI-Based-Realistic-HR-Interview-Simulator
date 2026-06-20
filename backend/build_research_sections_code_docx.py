from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT = Path(r"c:\Users\Public\ai-interview-platform")
OUT = ROOT / "backend" / "benchmark_outputs" / "research_sections_important_snippets_15_pages.docx"


def read_lines(rel_path: str):
    p = ROOT / rel_path
    return p.read_text(encoding="utf-8").splitlines()


def slice_lines(lines, start, end):
    return "\n".join(lines[start - 1:end])


def add_heading(doc: Document, text: str):
    p = doc.add_paragraph(text)
    r = p.runs[0]
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(12)


def add_code(doc: Document, code: str):
    p = doc.add_paragraph(code)
    for r in p.runs:
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(7.5)


SECTIONS = [
    ("3.1_four_model_specialized_question_generation", [
        ("backend/app/services/question_generation_service.py", 431, 530),
        ("backend/app/services/question_generation_service.py", 532, 573),
        ("backend/app/services/question_generation_service.py", 618, 647),
    ]),
    ("3.3_dual_round_rl_with_adaptive_difficulty", [
        ("backend/app/services/rl_adaptation_service.py", 50, 87),
        ("backend/app/services/rl_adaptation_service.py", 215, 260),
        ("backend/app/services/rl_adaptation_service.py", 357, 410),
    ]),
    ("3.4_real_time_multimodal_behavioral_analysis", [
        ("backend/app/services/multimodal_analysis_service.py", 76, 160),
        ("backend/app/services/multimodal_analysis_service.py", 703, 744),
        ("backend/app/services/multimodal_analysis_service.py", 841, 912),
    ]),
    ("3.5_two_phase_evaluation_architecture", [
        ("backend/app/services/ai_service.py", 504, 585),
        ("backend/app/services/ai_service.py", 693, 760),
        ("backend/app/services/ai_service.py", 835, 883),
        ("backend/app/utils/calibrator.py", 1, 43),
    ]),
    ("3.6_explainability_and_fairness_auditing", [
        ("backend/app/services/explainability_service.py", 207, 260),
        ("backend/app/services/fairness_service.py", 54, 145),
    ]),
    ("3.7_ai_proctoring_and_integrity_verification", [
        ("backend/app/services/proctoring_service.py", 97, 114),
        ("backend/app/services/proctoring_service.py", 623, 672),
        ("backend/app/services/proctoring_service.py", 853, 915),
        ("backend/app/routers/websocket.py", 31, 72),
    ]),
    ("3.8_practice_mode_and_personalized_development_roadmap", [
        ("backend/app/services/practice_mode_service.py", 203, 290),
        ("backend/app/services/development_roadmap_service.py", 64, 120),
    ]),
]


def main():
    doc = Document()

    for section_name, blocks in SECTIONS:
        add_heading(doc, section_name)
        for rel_path, start, end in blocks:
            add_heading(doc, f"# {rel_path}:{start}-{end}")
            lines = read_lines(rel_path)
            code = slice_lines(lines, start, min(end, len(lines)))
            add_code(doc, code)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(str(OUT))


if __name__ == "__main__":
    main()
